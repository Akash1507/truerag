import io

import docx
import pypdf

from app.core.errors import ParseError, PermanentIngestionError
from app.utils.observability import get_logger

logger = get_logger(__name__)

try:
    import pdfplumber
except ImportError:  # pragma: no cover - optional dependency during local dev
    pdfplumber = None

try:
    from pdf2image import convert_from_bytes
except ImportError:  # pragma: no cover - optional dependency during local dev
    convert_from_bytes = None

try:
    import pytesseract
except ImportError:  # pragma: no cover - optional dependency during local dev
    pytesseract = None


def _sanitize(text: str) -> str:
    """Remove lone surrogates produced by malformed PDF fonts."""
    return text.encode("utf-8", errors="ignore").decode("utf-8")


def parse_document(content: bytes, file_type: str) -> str:
    match file_type:
        case "txt" | "md":
            return _sanitize(_parse_text(content))
        case "pdf":
            return _sanitize(_parse_pdf(content))
        case "docx":
            return _sanitize(_parse_docx(content))
        case _:
            raise ParseError(f"Unsupported file type for pipeline parsing: {file_type}")


def _parse_text(content: bytes) -> str:
    try:
        text = content.decode("utf-8")
        if not text.strip():
            raise ParseError("Text document has no text")
        return text
    except Exception as e:
        if isinstance(e, ParseError):
            raise
        raise ParseError(f"Text decode failed: {e}") from e


def _parse_pdf(content: bytes) -> str:
    pymupdf_text = _parse_pdf_pymupdf(content) or ""
    pypdf_text = _parse_pdf_pypdf(content) or ""

    if len(pymupdf_text.strip()) < 50 and len(pypdf_text.strip()) < 50:
        text, page_count = _parse_pdf_ocr(content)
        logger.info(
            "ocr_fallback",
            extra={
                "operation": "parse_pdf",
                "extra_data": {
                    "page_count": page_count,
                    "char_count": len(text),
                },
            },
        )
        if not text.strip():
            raise ParseError("PDF OCR produced no text")
        return text

    use_pymupdf = len(pymupdf_text.strip()) >= len(pypdf_text.strip())
    text = pymupdf_text if use_pymupdf else pypdf_text
    if not text.strip():
        raise ParseError("PDF produced no text")

    if use_pymupdf:
        try:
            table_text = _extract_pdf_tables(content)
            if table_text.strip():
                text = f"{text}\n\n---TABLE---\n{table_text}"
        except Exception as exc:
            logger.warning(
                "pdf_table_extraction_failed",
                extra={
                    "operation": "parse_pdf",
                    "extra_data": {"error": str(exc)},
                },
            )

    return text


def _parse_pdf_ocr(content: bytes) -> tuple[str, int]:
    if convert_from_bytes is None or pytesseract is None:
        raise PermanentIngestionError("OCR dependencies are unavailable")
    try:
        images = convert_from_bytes(content, dpi=300)
        text = "\n\n".join(pytesseract.image_to_string(image) for image in images)
        return text, len(images)
    except pytesseract.TesseractNotFoundError as exc:
        raise PermanentIngestionError("tesseract-ocr binary not found") from exc
    except Exception as exc:
        raise ParseError(f"PDF OCR failed: {exc}") from exc


def _extract_pdf_tables(content: bytes) -> str:
    if pdfplumber is None:
        return ""

    table_texts: list[str] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            for table in page.extract_tables():
                if not table:
                    continue
                rows = [" | ".join(str(cell or "") for cell in row) for row in table if row]
                if not rows:
                    continue
                table_texts.append("\n".join(f"| {row} |" for row in rows))
    return "\n\n---TABLE---\n".join(table_texts)


def _parse_pdf_pymupdf(content: bytes) -> str | None:
    try:
        import fitz  # pymupdf
        doc = fitz.open(stream=content, filetype="pdf")
        pages = [page.get_text() for page in doc]
        doc.close()
        return "\n".join(pages)
    except Exception:
        return None


def _parse_pdf_pypdf(content: bytes) -> str | None:
    try:
        reader = pypdf.PdfReader(io.BytesIO(content))
        if not reader.pages:
            return None
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception:
        return None


def _parse_docx(content: bytes) -> str:
    try:
        doc = docx.Document(io.BytesIO(content))
        paragraph_text = "\n".join(p.text for p in doc.paragraphs)
        table_blocks: list[str] = []
        for table in doc.tables:
            rows = [" | ".join(cell.text for cell in row.cells) for row in table.rows]
            if not rows:
                continue
            table_blocks.append("\n".join(f"| {row} |" for row in rows))

        text = paragraph_text
        if table_blocks:
            table_text = "\n\n---TABLE---\n".join(table_blocks)
            if text.strip():
                text = f"{text}\n\n---TABLE---\n{table_text}"
            else:
                text = table_text

        if not text.strip():
            raise ParseError("DOCX produced no text")
        return text
    except Exception as exc:
        if isinstance(exc, ParseError):
            raise
        raise ParseError(f"DOCX parse failed: {exc}") from exc
